#!/usr/bin/env python3
"""文档格式生成脚本 — 只负责生成，内容理解由 Claude 原生 Read 工具完成"""

import argparse
import sys
import re
from pathlib import Path


def md_to_docx(md_path, output=None):
    """Markdown 转 Word（保留标题层级、列表、代码块、引用）"""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    src = Path(md_path)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    for line in src.read_text(encoding="utf-8").split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            heading = doc.add_heading(text, level=level)
            for run in heading.runs:
                run.font.name = "微软雅黑"
                if run.element.rPr is None:
                    run._r.get_or_add_rPr()
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        elif stripped.startswith("```"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:], style="Intense Quote")
        else:
            doc.add_paragraph(stripped)

    output = output or str(src.with_suffix(".docx"))
    doc.save(output)
    print(f"已转换: {output}")


def md_to_pdf(md_path, output=None):
    """Markdown 转 PDF（pandoc 渲染，保留完整格式）"""
    import subprocess

    src = Path(md_path)
    html_path = str(src.with_suffix(".html"))
    out_path = output or str(src.with_suffix(".pdf"))

    result = subprocess.run(
        ["pandoc", str(src), "-t", "html5", "--standalone",
         "--metadata", "lang=zh-CN", "-o", html_path],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print(f"pandoc 失败: {result.stderr}", file=sys.stderr)
        sys.exit(1)

    from weasyprint import HTML, CSS
    custom_css = CSS(string="""
        @font-face {
            font-family: 'CN';
            src: local('PingFang SC'), local('STHeiti'), local('Microsoft YaHei');
        }
        body { font-family: 'CN', sans-serif; font-size: 12pt; line-height: 1.8; padding: 20mm; }
        h1 { font-size: 22pt; border-bottom: 2px solid #333; }
        h2 { font-size: 18pt; border-bottom: 1px solid #999; }
        h3 { font-size: 15pt; }
        pre { background: #f5f5f5; padding: 10px; overflow: auto; }
        code { font-family: 'Courier New', monospace; font-size: 10pt; }
        table { border-collapse: collapse; width: 100%; margin: 10px 0; }
        th, td { border: 1px solid #999; padding: 6px; }
        th { background: #f0f0f0; font-weight: bold; }
        blockquote { border-left: 3px solid #ccc; padding-left: 10px; color: #666; }
    """)
    HTML(filename=html_path).write_pdf(out_path, stylesheets=[custom_css])
    Path(html_path).unlink(missing_ok=True)
    print(f"已转换: {out_path}")


def docx_to_pdf(docx_path, output=None):
    """Word 转 PDF（mammoth 提取 HTML → weasyprint 渲染）"""
    import subprocess

    src = Path(docx_path)
    html_path = str(src.with_suffix(".html"))
    out_path = output or str(src.with_suffix(".pdf"))

    try:
        import mammoth
        with open(src, "rb") as f:
            result = mammoth.convert_to_html(f)
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(result.value)
    except ImportError:
        subprocess.run(
            ["pandoc", str(src), "-t", "html5", "--standalone", "-o", html_path],
            check=True
        )

    from weasyprint import HTML, CSS
    custom_css = CSS(string="""
        @font-face {
            font-family: 'CN';
            src: local('PingFang SC'), local('STHeiti'), local('Microsoft YaHei');
        }
        body { font-family: 'CN', sans-serif; font-size: 12pt; line-height: 1.8; padding: 20mm; }
        h1 { font-size: 22pt; border-bottom: 2px solid #333; }
        h2 { font-size: 18pt; border-bottom: 1px solid #999; }
        h3 { font-size: 15pt; }
        table { border-collapse: collapse; width: 100%; margin: 10px 0; }
        th, td { border: 1px solid #999; padding: 6px; }
        th { background: #f0f0f0; font-weight: bold; }
        ul, ol { padding-left: 20px; }
    """)
    HTML(filename=html_path).write_pdf(out_path, stylesheets=[custom_css])
    Path(html_path).unlink(missing_ok=True)
    print(f"已转换: {out_path}")


def merge(file_list, output="merged_output.docx"):
    """合并多个 Word 文档"""
    from docx import Document
    merged = Document()
    for filepath in file_list:
        src = Document(filepath)
        for para in src.paragraphs:
            merged.add_paragraph(para.text)
        merged.add_page_break()
    merged.save(output)
    print(f"已合并为: {output}")


def extract(filepath, output_type="text"):
    """提取文档内容 — 实际由 Claude Read 工具完成，此处保留接口"""
    print(f"[提示: 请直接使用 Read 工具读取 {filepath}，Claude 会自动理解内容]")


def summary(filepath):
    """文档摘要 — 由 Claude 原生完成"""
    print(f"[提示: 请使用 Read 工具读取 {filepath}，然后让 Claude 生成摘要]")


def review(filepath):
    """文档质检 — 由 Claude 原生完成"""
    print(f"[提示: 请使用 Read 工具读取 {filepath}，然后让 Claude 进行质检]")


def main():
    parser = argparse.ArgumentParser(description="文档格式生成工具（内容理解由 Claude Read 工具完成）")
    subparsers = parser.add_subparsers(dest="operation", required=True)

    # md2docx
    p_md = subparsers.add_parser("md2docx")
    p_md.add_argument("file", help="Markdown 文件路径")
    p_md.add_argument("--output")

    # md2pdf
    p_md2pdf = subparsers.add_parser("md2pdf")
    p_md2pdf.add_argument("file", help="Markdown 文件路径")
    p_md2pdf.add_argument("--output")

    # docx2pdf
    p_docx2pdf = subparsers.add_parser("docx2pdf")
    p_docx2pdf.add_argument("file", help="Word 文件路径")
    p_docx2pdf.add_argument("--output")

    # merge
    p_merge = subparsers.add_parser("merge")
    p_merge.add_argument("files", nargs="+")
    p_merge.add_argument("--output", default="merged_output.docx")

    # extract（占位，实际由 Read 工具完成）
    p_extract = subparsers.add_parser("extract")
    p_extract.add_argument("file")
    p_extract.add_argument("--type", default="text", choices=["text", "outline"])

    # summary（占位，实际由 Claude 完成）
    p_summary = subparsers.add_parser("summary")
    p_summary.add_argument("file")

    # review（占位，实际由 Claude 完成）
    p_review = subparsers.add_parser("review")
    p_review.add_argument("file")

    args = parser.parse_args()

    if args.operation == "md2docx":
        md_to_docx(args.file, args.output)
    elif args.operation == "md2pdf":
        md_to_pdf(args.file, args.output)
    elif args.operation == "docx2pdf":
        docx_to_pdf(args.file, args.output)
    elif args.operation == "merge":
        merge(args.files, args.output)
    elif args.operation == "extract":
        extract(args.file, args.type)
    elif args.operation == "summary":
        summary(args.file)
    elif args.operation == "review":
        review(args.file)


if __name__ == "__main__":
    main()
